import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.tree import DecisionTreeClassifier
from sklearn.metrics import accuracy_score, classification_report
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

# Verileri oku
df = pd.read_csv("algerian-bejaia.csv")

# Hedef değişken ve özellikler
X = df.drop(columns=["Fire Classes"])
y = df["Fire Classes"]

# Verileri eğitim ve test setlerine ayır
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Karar ağacı modelini oluştur
clf = DecisionTreeClassifier(random_state=42)

# Modeli eğit
clf.fit(X_train, y_train)

# Test seti üzerinde tahmin yap
y_pred = clf.predict(X_test)

# Doğruluk oranını hesapla
accuracy = accuracy_score(y_test, y_pred)

# Model performansını değerlendir (classification report)
class_report = classification_report(y_test, y_pred, output_dict=True)

# Classification Report'u DataFrame'e çevir
df_report = pd.DataFrame(class_report).transpose()

# Word belgesi oluştur
doc = Document()
doc.add_heading('Classification Report', 0)

# Doğruluk oranını belgeye ekle
doc.add_paragraph(f"Accuracy: {accuracy:.6f}")

# Tabloyu Word belgesine ekle
table = doc.add_table(df_report.shape[0] + 2, df_report.shape[1] + 1)  # +2: Bir satır accuracy için, bir satır başlık için
table.style = 'Table Grid'

# Accuracy başlığını ekle
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metrics'
hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
hdr_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

# Diğer başlıkları ekle
for i, col in enumerate(df_report.columns):
    hdr_cells[i + 1].text = col
    hdr_cells[i + 1].paragraphs[0].runs[0].font.size = Pt(12)
    hdr_cells[i + 1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

# Verileri ekle
for i in range(df_report.shape[0]):
    row_cells = table.rows[i + 1].cells
    row_cells[0].text = df_report.index[i]  # Metriklerin isimlerini ekleyin
    for j in range(df_report.shape[1]):
        row_cells[j + 1].text = f"{df_report.values[i, j]:.6f}"  # 6 basamaklı virgülden sonra formatla
        row_cells[j + 1].paragraphs[0].runs[0].font.size = Pt(10)

# Accuracy değerini ekleyin
row_accuracy = table.rows[-1].cells
row_accuracy[0].text = 'Accuracy'
row_accuracy[1].text = f"{accuracy:.6f}"
row_accuracy[1].paragraphs[0].runs[0].font.size = Pt(10)

# Word belgesini kaydet
doc.save('classification_report.docx')
print("Classification Report başarıyla 'classification_report.docx' dosyasına kaydedildi.")
